import os
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from docx import Document
from sentence_transformers import SentenceTransformer, util
import re
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from tqdm import tqdm
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
import datetime
import time

# NLTK resurslarini yuklash
nltk.download('punkt', quiet=True)
nltk.download('stopwords', quiet=True)

# Word fayldan matn chiqarish
def extract_text_from_docx(docx_path):
    try:
        doc = Document(docx_path)
        full_text = []
        
        # Paragraflarni olish
        for para in doc.paragraphs:
            if para.text.strip():  # Bo'sh paragraflarni o'tkazib yuborish
                full_text.append(para.text)
        
        # Jadvallardan matnni olish
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        return "\n".join(full_text)
    except Exception as e:
        print(f"Xato: {docx_path} faylini o'qishda muammo: {str(e)}")
        return ""

# Matnni tozalash 
def clean_text(text):
    # Kichik harflarga o'girish
    text = text.lower()
    
    # Belgilarni olib tashlash (harflar va sonlar qoladi)
    text = re.sub(r'[^\w\s]', ' ', text)
    
    # Ko'p bo'shliqlarni bitta bo'shliqga almashtirish
    text = re.sub(r'\s+', ' ', text)
    
    return text.strip()

# Stop so'zlarni olib tashlash
def remove_stopwords(text, language='uzbek'):
    if language == 'uzbek':
        # Uzbek tilidagi stop so'zlar (misol uchun)
        uzb_stopwords = {'va', 'ham', 'bilan', 'bu', 'shu', 'uchun', 'bo\'yicha', 'kerak', 'lekin', 'ammo', 
                         'chunki', 'agar', 'yoki', 'bir', 'har', 'qanday', 'hamma', 'barcha'}
        words = word_tokenize(text)
        return ' '.join([word for word in words if word not in uzb_stopwords])
    else:
        # Ingliz tilidagi stop so'zlar
        stop_words = set(stopwords.words('english'))
        words = word_tokenize(text)
        return ' '.join([word for word in words if word not in stop_words])

# TF-IDF asosida o'xshashlikni tekshirish
def check_plagiarism_tfidf(text1, text2):
    # TF-IDF vektorlashtirish
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform([text1, text2])
    
    # Kosinusli o'xshashlikni hisoblash
    similarity = (tfidf_matrix * tfidf_matrix.T).toarray()[0, 1]
    return similarity

# Neyron tarmoq yordamida semantik o'xshashlikni aniqlash
def check_plagiarism_neural(text1, text2, model, threshold=0.8):
    # Umumiy matnlarni vektorlashtirish
    embedding1 = model.encode(text1, convert_to_tensor=True)
    embedding2 = model.encode(text2, convert_to_tensor=True)
    
    # Umumiy o'xshashlikni hisoblash (Cosine Similarity)
    overall_similarity = util.cos_sim(embedding1, embedding2)[0][0].item()
    
    return overall_similarity

# Paragraf darajasida o'xshashlikni tekshirish
def check_paragraph_similarity(paragraphs1, paragraphs2, model, threshold=0.8):
    similar_paragraphs = []
    
    # Paragraflarni olish (bo'sh qatorlarni ajratuvchi sifatida)
    if isinstance(paragraphs1, str):
        paragraphs1 = paragraphs1.split('\n')
    if isinstance(paragraphs2, str):
        paragraphs2 = paragraphs2.split('\n')
    
    # Faqat ma'lumotli paragraflarni qoldirish
    paragraphs1 = [p.strip() for p in paragraphs1 if len(p.strip()) > 30]
    paragraphs2 = [p.strip() for p in paragraphs2 if len(p.strip()) > 30]
    
    # Har bir paragrafni tekshirish
    for i, p1 in enumerate(paragraphs1):
        max_sim = 0
        best_match = ""
        best_idx = -1
        
        for j, p2 in enumerate(paragraphs2):
            # Paragraflarni vektorlashtirish
            emb1 = model.encode(p1, convert_to_tensor=True)
            emb2 = model.encode(p2, convert_to_tensor=True)
            
            # O'xshashlikni hisoblash
            sim = util.cos_sim(emb1, emb2)[0][0].item()
            
            if sim > max_sim:
                max_sim = sim
                best_match = p2
                best_idx = j
        
        # Agar o'xshashlik darajasi yuqori bo'lsa
        if max_sim > threshold:
            similar_paragraphs.append({
                "doc1_paragraph": p1,
                "doc2_paragraph": best_match,
                "doc1_index": i,
                "doc2_index": best_idx,
                "similarity": max_sim
            })
    
    return similar_paragraphs

# Jumlalar o'xshashligini tekshirish
def check_sentence_similarity(text1, text2, model, threshold=0.8):
    # Matnni jumlaarga bo'lish
    sentences1 = sent_tokenize(text1)
    sentences2 = sent_tokenize(text2)
    
    similar_sentences = []
    
    # Progress bar yaratish
    total_comparisons = len(sentences1) * len(sentences2)
    with tqdm(total=total_comparisons, desc="Jumlalarni tekshirish") as pbar:
        # Har bir jumlani tekshirish
        for i, s1 in enumerate(sentences1):
            if len(s1.split()) < 5:  # Juda qisqa jumlalarni o'tkazib yuborish
                continue
                
            best_match = {"sentence": "", "similarity": 0, "index": -1}
            
            for j, s2 in enumerate(sentences2):
                if len(s2.split()) < 5:  # Juda qisqa jumlalarni o'tkazib yuborish
                    continue
                    
                # Jumlalarni vektorlashtirish
                emb1 = model.encode(s1, convert_to_tensor=True)
                emb2 = model.encode(s2, convert_to_tensor=True)
                
                # O'xshashlikni hisoblash
                sim = util.cos_sim(emb1, emb2)[0][0].item()
                
                if sim > best_match["similarity"]:
                    best_match = {"sentence": s2, "similarity": sim, "index": j}
                
                pbar.update(1)
            
            # Agar o'xshashlik darajasi yuqori bo'lsa
            if best_match["similarity"] > threshold:
                similar_sentences.append({
                    "doc1_sentence": s1,
                    "doc2_sentence": best_match["sentence"],
                    "doc1_index": i,
                    "doc2_index": best_match["index"],
                    "similarity": best_match["similarity"]
                })
    
    return similar_sentences

# Natijalarni vizualizatsiya qilish
def visualize_similarity(similarity, paragraph_sim, sentence_sim, output_dir="./plagiarism_report"):
    # Hisobot papkasini yaratish
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # O'xshashlik darajalari diagrammasi
    labels = ['Umumiy', 'TF-IDF', 'Semantik']
    values = [similarity['overall'] * 100, similarity['tfidf'] * 100, similarity['neural'] * 100]
    
    plt.figure(figsize=(10, 6))
    bars = plt.bar(labels, values, color=['blue', 'green', 'red'])
    plt.axhline(y=70, color='r', linestyle='--', alpha=0.7)  # Plagiat chegarasi
    
    # Bar ustiga foizlarni yozish
    for bar in bars:
        height = bar.get_height()
        plt.text(bar.get_x() + bar.get_width()/2., height + 1,
                 f'{height:.1f}%', ha='center', va='bottom')
    
    plt.title('O\'xshashlik darajalari')
    plt.ylim(0, 100)
    plt.ylabel('O\'xshashlik foizi (%)')
    plt.savefig(f"{output_dir}/similarity_chart_{timestamp}.png", dpi=300, bbox_inches='tight')
    
    # Paragraf va jumlalardagi o'xshashlik diagrammasi
    if paragraph_sim:
        paragraph_similarities = [item['similarity'] * 100 for item in paragraph_sim]
        plt.figure(figsize=(12, 6))
        plt.hist(paragraph_similarities, bins=10, alpha=0.7, color='blue')
        plt.title('Paragraflar o\'xshashligi')
        plt.xlabel('O\'xshashlik foizi (%)')
        plt.ylabel('Paragraflar soni')
        plt.xlim(0, 100)
        plt.grid(alpha=0.3)
        plt.savefig(f"{output_dir}/paragraph_similarity_{timestamp}.png", dpi=300, bbox_inches='tight')
    
    if sentence_sim:
        sentence_similarities = [item['similarity'] * 100 for item in sentence_sim]
        plt.figure(figsize=(12, 6))
        plt.hist(sentence_similarities, bins=10, alpha=0.7, color='green')
        plt.title('Jumlalar o\'xshashligi')
        plt.xlabel('O\'xshashlik foizi (%)')
        plt.ylabel('Jumlalar soni')
        plt.xlim(0, 100)
        plt.grid(alpha=0.3)
        plt.savefig(f"{output_dir}/sentence_similarity_{timestamp}.png", dpi=300, bbox_inches='tight')
    
    # Heatmap uchun ma'lumotlar
    if paragraph_sim and len(paragraph_sim) > 0:
        plt.figure(figsize=(12, 8))
        heatmap_data = np.zeros((min(len(paragraph_sim), 20), 4))
        
        for i, item in enumerate(paragraph_sim[:20]):  # Faqat birinchi 20 natija
            heatmap_data[i, 0] = item['doc1_index']
            heatmap_data[i, 1] = item['doc2_index']
            heatmap_data[i, 2] = item['similarity'] * 100
            heatmap_data[i, 3] = len(item['doc1_paragraph'].split())
        
        df = pd.DataFrame(heatmap_data, columns=['Doc1 Index', 'Doc2 Index', 'Similarity %', 'Word Count'])
        
        sns.heatmap(df.pivot(index='Doc1 Index', columns='Doc2 Index', values='Similarity %'), 
                   cmap='YlOrRd', annot=True, fmt='.1f')
        plt.title('Paragraflar orasidagi o\'xshashlik heatmap')
        plt.savefig(f"{output_dir}/paragraph_heatmap_{timestamp}.png", dpi=300, bbox_inches='tight')
    
    plt.close('all')
    return f"{output_dir}/similarity_chart_{timestamp}.png"

# Hisobot yaratish
def generate_report(doc1_name, doc2_name, similarity, paragraph_sim, sentence_sim, output_dir="./plagiarism_report"):
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    report_file = f"{output_dir}/plagiarism_report_{timestamp}.html"
    
    # O'xshash paragraf va jumlalar jadvalini yaratish
    paragraph_table = ""
    if paragraph_sim:
        paragraph_table = "<h2>O'xshash paragraflar:</h2>"
        paragraph_table += "<table border='1' style='width:100%; border-collapse: collapse;'>"
        paragraph_table += "<tr><th>№</th><th>1-hujjat paragrafi</th><th>2-hujjat paragrafi</th><th>O'xshashlik</th></tr>"
        
        for i, item in enumerate(paragraph_sim):
            paragraph_table += f"<tr><td>{i+1}</td><td>{item['doc1_paragraph']}</td><td>{item['doc2_paragraph']}</td><td>{item['similarity']*100:.1f}%</td></tr>"
        
        paragraph_table += "</table>"
    
    sentence_table = ""
    if sentence_sim:
        sentence_table = "<h2>O'xshash jumlalar:</h2>"
        sentence_table += "<table border='1' style='width:100%; border-collapse: collapse;'>"
        sentence_table += "<tr><th>№</th><th>1-hujjat jumlasi</th><th>2-hujjat jumlasi</th><th>O'xshashlik</th></tr>"
        
        for i, item in enumerate(sentence_sim):
            sentence_table += f"<tr><td>{i+1}</td><td>{item['doc1_sentence']}</td><td>{item['doc2_sentence']}</td><td>{item['similarity']*100:.1f}%</td></tr>"
        
        sentence_table += "</table>"
    
    # O'xshashlik darajasi tasnifi
    plagiarism_level = "Past"
    if similarity['overall'] > 0.7:
        plagiarism_level = "Juda yuqori"
    elif similarity['overall'] > 0.5:
        plagiarism_level = "Yuqori"
    elif similarity['overall'] > 0.3:
        plagiarism_level = "O'rta"
    
    # Hisobot HTML shaklida
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>Plagiat tekshiruvi hisoboti</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 20px; }}
            h1 {{ color: #2c3e50; }}
            h2 {{ color: #3498db; margin-top: 20px; }}
            .summary {{ background-color: #f9f9f9; padding: 15px; border-radius: 5px; margin: 15px 0; }}
            .warning {{ background-color: #ffe0e0; padding: 10px; border-radius: 5px; margin: 15px 0; }}
            table {{ border-collapse: collapse; width: 100%; margin-bottom: 20px; }}
            th, td {{ padding: 8px; text-align: left; border: 1px solid #ddd; }}
            th {{ background-color: #f2f2f2; }}
            tr:nth-child(even) {{ background-color: #f9f9f9; }}
            .progress-container {{ width: 100%; background-color: #e0e0e0; border-radius: 5px; }}
            .progress-bar {{ height: 24px; border-radius: 5px; }}
        </style>
    </head>
    <body>
        <h1>Plagiat tekshiruvi hisoboti</h1>
        <p><strong>Tekshirilgan sana:</strong> {datetime.datetime.now().strftime("%Y-%m-%d %H:%M")}</p>
        <p><strong>1-hujjat nomi:</strong> {doc1_name}</p>
        <p><strong>2-hujjat nomi:</strong> {doc2_name}</p>
        
        <div class="summary">
            <h2>Umumiy natijalar:</h2>
            <p><strong>Umumiy o'xshashlik darajasi:</strong> {similarity['overall']*100:.2f}%</p>
            <p><strong>TF-IDF o'xshashlik:</strong> {similarity['tfidf']*100:.2f}%</p>
            <p><strong>Semantik o'xshashlik:</strong> {similarity['neural']*100:.2f}%</p>
            <p><strong>Plagiat darajasi:</strong> {plagiarism_level}</p>
            
            <div class="progress-container">
                <div class="progress-bar" style="width: {similarity['overall']*100}%; background-color: {
                    '#ff4d4d' if similarity['overall'] > 0.7 else 
                    '#ffa64d' if similarity['overall'] > 0.5 else 
                    '#ffff4d' if similarity['overall'] > 0.3 else 
                    '#4dff4d'
                };">{similarity['overall']*100:.1f}%</div>
            </div>
        </div>
        
        {
        '<div class="warning"><strong>Ogohlantirish:</strong> Ushbu ikki hujjat orasida sezilarli darajada o\'xshashlik aniqlandi.</div>' if similarity['overall'] > 0.5 else ''
        }
        
        {paragraph_table}
        
        {sentence_table}
        
        <p><em>Eslatma: Ushbu hisobot avtomatik tahlil natijasi hisoblanadi va mutlaq haqiqat emas. Yakuniy qaror qabul qilish uchun mutaxassis ko'rigi tavsiya etiladi.</em></p>
    </body>
    </html>
    """
    
    with open(report_file, "w", encoding="utf-8") as f:
        f.write(html_content)
    
    return report_file

# Asosiy funksiya: diplom ishlarini tekshirish
def check_diplom_plagiarism(doc1_path, doc2_path, threshold=0.7, language='uzbek', output_dir="./plagiarism_report"):
    start_time = time.time()
    print(f"Tahlil boshlandi... Iltimos kuting.")
    
    # Matnlarni chiqarish
    doc1_text = extract_text_from_docx(doc1_path)
    doc2_text = extract_text_from_docx(doc2_path)
    
    if not doc1_text or not doc2_text:
        print("Xato: Fayllardan matn o'qib bo'lmadi. Fayl mavjudligini va to'g'ri formatda ekanligini tekshiring.")
        return
    
    print(f"Fayllar muvaffaqiyatli o'qildi. Matn tahlili boshlandi...")
    
    # Matnlarni tozalash
    doc1_clean = clean_text(doc1_text)
    doc2_clean = clean_text(doc2_text)
    
    # Stop so'zlarni olib tashlash
    doc1_no_stop = remove_stopwords(doc1_clean, language)
    doc2_no_stop = remove_stopwords(doc2_clean, language)
    
    # TF-IDF asosida o'xshashlikni tekshirish
    tfidf_similarity = check_plagiarism_tfidf(doc1_no_stop, doc2_no_stop)
    
    # Sentence Transformers modelini yuklash
    print("Neyron tarmoq modelini yuklash...")
    model = SentenceTransformer('all-MiniLM-L6-v2')
    
    # Semantik o'xshashlikni tekshirish
    neural_similarity = check_plagiarism_neural(doc1_text, doc2_text, model, threshold)
    
    # Umumiy o'xshashlik - ikki usul o'rtacha qiymati
    overall_similarity = (tfidf_similarity + neural_similarity) / 2
    
    # O'xshash paragraflarni topish
    print("Paragraflar o'xshashligini tekshirish...")
    similar_paragraphs = check_paragraph_similarity(doc1_text, doc2_text, model, threshold)
    
    # O'xshash jumlalarni topish
    print("Jumlalar o'xshashligini tekshirish...")
    similar_sentences = check_sentence_similarity(doc1_text, doc2_text, model, threshold)
    
    # Natijalarni vizualizatsiya qilish
    print("Natijalarni tayyorlash...")
    similarity_data = {
        'overall': overall_similarity,
        'tfidf': tfidf_similarity,
        'neural': neural_similarity
    }
    
    # Hisobot yaratish
    doc1_name = os.path.basename(doc1_path)
    doc2_name = os.path.basename(doc2_path)
    
    chart_path = visualize_similarity(similarity_data, similar_paragraphs, similar_sentences, output_dir)
    report_path = generate_report(doc1_name, doc2_name, similarity_data, similar_paragraphs, similar_sentences, output_dir)
    
    # Natijalarni chop etish
    elapsed_time = time.time() - start_time
    print(f"\n=== PLAGIAT TEKSHIRUV NATIJALARI ===")
    print(f"Tahlil vaqti: {elapsed_time:.2f} sekund")
    print(f"TF-IDF o'xshashlik: {tfidf_similarity * 100:.2f}%")
    print(f"Semantik o'xshashlik: {neural_similarity * 100:.2f}%")
    print(f"Umumiy o'xshashlik darajasi: {overall_similarity * 100:.2f}%")
    
    # O'xshashlik darajasini talqin qilish
    if overall_similarity > 0.7:
        print("\nXULOSA: Juda yuqori darajadagi o'xshashlik aniqlandi. Bu plagiat bo'lishi mumkin.")
    elif overall_similarity > 0.5:
        print("\nXULOSA: Yuqori darajadagi o'xshashlik aniqlandi. Ko'rib chiqish tavsiya etiladi.")
    elif overall_similarity > 0.3:
        print("\nXULOSA: O'rta darajadagi o'xshashlik aniqlandi. Ayrim qismlar o'xshash bo'lishi mumkin.")
    else:
        print("\nXULOSA: Past darajadagi o'xshashlik aniqlandi. Hujjatlar asosan mustaqil.")
    
    print(f"\nO'xshash paragraflar soni: {len(similar_paragraphs)}")
    print(f"O'xshash jumlalar soni: {len(similar_sentences)}")
    
    if similar_paragraphs or similar_sentences:
        print(f"\nBatafsil hisobot yaratildi: {report_path}")
        print(f"Vizualizatsiya: {chart_path}")
    
    return {
        'overall_similarity': overall_similarity,
        'tfidf_similarity': tfidf_similarity,
        'neural_similarity': neural_similarity,
        'similar_paragraphs': similar_paragraphs,
        'similar_sentences': similar_sentences,
        'report_path': report_path,
        'chart_path': chart_path
    }

# Agar ushbu skript to'g'ridan-to'g'ri ishga tushirilsa
if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description='Diplom ishlari o\'xshashligini tekshirish tizimi')
    parser.add_argument('--doc1', required=True, help='Birinchi diplom ishi (DOCX formati)')
    parser.add_argument('--doc2', required=True, help='Ikkinchi diplom ishi (DOCX formati)')
    parser.add_argument('--threshold', type=float, default=0.7, help='O\'xshashlik chegarasi (0-1 oralig\'ida)')
    parser.add_argument('--language', choices=['uzbek', 'english'], default='uzbek', help='Matn tili')
    parser.add_argument('--output', default='./plagiarism_report', help='Hisobot papkasi')
    
    args = parser.parse_args()
    
    check_diplom_plagiarism(args.doc1, args.doc2, args.threshold, args.language, args.output)