#utils.py
import os
import re
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from docx import Document
import datetime
import time

# NLTK resurslarini yuklash
nltk.download('punkt', quiet=True)
nltk.download('stopwords', quiet=True)

# Word fayldan matn chiqarish
def extract_text_from_docx(docx_path):
    try:
        doc = Document(docx_path)
        full_text = []
        
        # Paragraflarni olish
        for para in doc.paragraphs:
            if para.text.strip():  # Bo'sh paragraflarni o'tkazib yuborish
                full_text.append(para.text)
        
        # Jadvallardan matnni olish
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        return full_text
    except Exception as e:
        print(f"Xato: {docx_path} faylini o'qishda muammo: {str(e)}")
        return []

# Matnni tozalash 
def clean_text(text):
    # Kichik harflarga o'girish
    text = text.lower()
    
    # Belgilarni olib tashlash (harflar va sonlar qoladi)
    text = re.sub(r'[^\w\s]', ' ', text)
    
    # Ko'p bo'shliqlarni bitta bo'shliqqa almashtirish
    text = re.sub(r'\s+', ' ', text)
    
    return text.strip()

# Stop so'zlarni olib tashlash
def remove_stopwords(text, language='uzbek'):
    if language == 'uzbek':
        # Uzbek tilidagi stop so'zlar (misol uchun)
        uzb_stopwords = {'va', 'ham', 'bilan', 'bu', 'shu', 'uchun', 'bo\'yicha', 'kerak', 'lekin', 'ammo', 
                         'chunki', 'agar', 'yoki', 'bir', 'har', 'qanday', 'hamma', 'barcha'}
        words = word_tokenize(text)
        return ' '.join([word for word in words if word not in uzb_stopwords])
    else:
        # Ingliz tilidagi stop so'zlar
        stop_words = set(stopwords.words('english'))
        words = word_tokenize(text)
        return ' '.join([word for word in words if word not in stop_words])

# Matnni boblar bo'yicha ajratish
def split_by_chapters(paragraphs):
    chapters = {}
    current_chapter = None
    chapter_pattern = re.compile(r"^(?:BOB|ГЛАВА|CHAPTER)\s+(\d+)", re.IGNORECASE)
    
    # Birinchi bo'lim (kirish qismi yoki mazmun)
    introduction_text = []
    
    for para in paragraphs:
        # Bob sarlavhasini aniqlash
        match = chapter_pattern.match(para.strip())
        
        if match:
            chapter_num = match.group(1)
            current_chapter = chapter_num
            chapters[current_chapter] = [para]
        elif current_chapter:
            chapters[current_chapter].append(para)
        else:
            introduction_text.append(para)
    
    # Kirish qismini ham qo'shish
    if introduction_text:
        chapters["0"] = introduction_text
    
    return chapters

# Time formatter
def format_elapsed_time(seconds):
    if seconds < 60:
        return f"{seconds:.2f} sekund"
    elif seconds < 3600:
        minutes = seconds // 60
        seconds = seconds % 60
        return f"{int(minutes)} daqiqa {int(seconds)} sekund"
    else:
        hours = seconds // 3600
        minutes = (seconds % 3600) // 60
        seconds = seconds % 60
        return f"{int(hours)} soat {int(minutes)} daqiqa {int(seconds)} sekund"